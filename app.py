import streamlit as st
import os
from dotenv import load_dotenv
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
import google.generativeai as genai
from io import BytesIO
import json
from datetime import datetime
import posixpath
from lxml import etree

# --- 1. FUNGSI TEKNIS MANIPULASI WORD (OXML) ---

def get_or_create_comments_part(doc):
    """Mendapatkan atau membuat part comments.xml dengan inisialisasi elemen langsung."""
    main_doc_part = doc.part
    
    # 1. Cek relasi yang sudah ada
    for rel in main_doc_part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return rel.target_part

    # 2. Buat elemen root secara langsung sebagai objek (bukan string parsing)
    # Ini menghindari error "Document is empty"
    comments_xml = OxmlElement('w:comments')
    
    part_dir = posixpath.dirname(main_doc_part.partname)
    comments_part_name = PackURI(posixpath.join(part_dir, 'comments.xml'))
    
    # 3. Buat part dengan objek XML langsung
    comments_part = Part(
        comments_part_name,
        CT.WML_COMMENTS,
        comments_xml, # Mengirim objek elemen, bukan string
        main_doc_part.package
    )
    main_doc_part.relate_to(comments_part, RT.COMMENTS)
    return comments_part

def add_comment_to_paragraph(paragraph, comment_text, author="Author"):
    """Menyisipkan balon komentar dengan logika pemisahan objek high-level dan low-level XML."""
    
    try:
        if not paragraph.text.strip():
            st.error("Gagal menyisipkan komentar: Paragraph kosong.")
            return
        # 1. Akses part komentar
        doc = paragraph.part.package.main_document_part.document
        comments_part = get_or_create_comments_part(doc)
        
        # 2. Ambil root XML comments dengan error handling
        comments_xml = OxmlElement('w:comments')  # Inisialisasi default
        try:
            if hasattr(comments_part, 'element') and comments_part.element is not None:
                comments_xml = comments_part.element
            elif hasattr(comments_part, '_element') and comments_part._element is not None:
                comments_xml = comments_part._element
            elif hasattr(comments_part, 'blob') and comments_part.blob is not None and len(comments_part.blob) > 0:
                comments_xml = parse_xml(comments_part.blob)
        except Exception as parse_error:
            # Jika parsing gagal, gunakan elemen default yang sudah dibuat
            st.warning(f"XML comments tidak valid, membuat baru. Error: {str(parse_error)[:50]}")
            comments_xml = OxmlElement('w:comments')
        
        # 3. Buat ID unik
        existing_comments = comments_xml.xpath('//w:comment')
        comment_id = str(len(existing_comments) + 1)
        
        # 4. Buat elemen komentar dengan struktur lengkap
        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), comment_id)
        comment.set(qn('w:author'), author)
        comment.set(qn('w:date'), datetime.now().isoformat())
        comment.set(qn('w:initials'), author[0] if author else 'A')
        
        # Buat paragraph dalam komentar dengan properties
        p_comm = OxmlElement('w:p')
        
        # Tambahkan paragraph properties
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'CommentReference')
        pPr.append(pStyle)
        p_comm.append(pPr)
        
        # Tambahkan reference run
        r_ref = OxmlElement('w:r')
        rPr_ref = OxmlElement('w:rPr')
        rStyle_ref = OxmlElement('w:rStyle')
        rStyle_ref.set(qn('w:val'), 'CommentReference')
        rPr_ref.append(rStyle_ref)
        r_ref.append(rPr_ref)
        annotationRef = OxmlElement('w:annotationRef')
        r_ref.append(annotationRef)
        p_comm.append(r_ref)
        
        # Tambahkan text run dengan konten komentar
        r_text = OxmlElement('w:r')
        t_comm = OxmlElement('w:t')
        t_comm.set(qn('xml:space'), 'preserve')
        t_comm.text = comment_text
        r_text.append(t_comm)
        p_comm.append(r_text)
        
        comment.append(p_comm)
        comments_xml.append(comment)
        
        # PENTING: Serialize comments_xml ke blob agar benar-benar tersimpan
        try:
            # Serialize dengan proper XML declaration
            xml_bytes = etree.tostring(
                comments_xml,
                xml_declaration=True,
                encoding='UTF-8',
                standalone=True
            )
            # Simpan ke _blob comments_part
            comments_part._blob = xml_bytes
        except Exception as serialize_error:
            st.warning(f"Error saat menyimpan XML: {str(serialize_error)[:50]}")
        
        # Alternative: Update element juga untuk consistency
        if hasattr(comments_part, '_element'):
            comments_part._element = comments_xml
        
        # 5. Injeksi Range ke Dokumen Utama
        # Menggunakan paragraph._p untuk manipulasi XML mentah (RangeStart & End)
        p_element = paragraph._p
        
        start = OxmlElement('w:commentRangeStart')
        start.set(qn('w:id'), comment_id)
        p_element.insert(0, start)
        
        end = OxmlElement('w:commentRangeEnd')
        end.set(qn('w:id'), comment_id)
        p_element.append(end)
        
        # 6. Menambahkan Referensi (Anchor)
        # PERBAIKAN: Gunakan objek paragraph (bukan _p) untuk add_run()
        new_run = paragraph.add_run() 
        ref = OxmlElement('w:commentReference')
        ref.set(qn('w:id'), comment_id)
        new_run._r.append(ref)
        
    except Exception as e:
        st.error(f"Gagal menyisipkan komentar: {e}")

# --- 2. FUNGSI ALUR KERJA ---

def find_section_paragraph(doc, keywords):
    """Cari paragraph section dengan exact match terlebih dahulu, 
    kemudian fallback ke case-insensitive jika tidak ketemu.
    """
    # Pass 1: Exact match (case-sensitive)
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        
        if not para_text:
            continue
        
        # Exact match untuk keywords
        for kw in keywords:
            if para_text == kw:  # Exact match, case-sensitive
                return para_idx
    
    # Pass 2: Case-insensitive match sebagai fallback
    for para_idx, para in enumerate(doc.paragraphs):
        para_text_lower = para.text.strip().lower()
        
        if not para_text_lower:
            continue
        
        for kw in keywords:
            if para_text_lower == kw.lower():  # Case-insensitive match
                return para_idx
    
    return None

def process_word_with_comments(uploaded_file, review_data):
    """Mencari bab yang sesuai dan menyisipkan komentar terpilih."""
    file_bytes = uploaded_file.getvalue()
    doc = Document(BytesIO(file_bytes))
    # Cek dokumen kosong
    if not doc.paragraphs or all(not para.text.strip() for para in doc.paragraphs):
        st.error("Dokumen Word kosong, tidak bisa menyisipkan komentar.")
        return BytesIO()
    
    sections = {
        "Judul": ["JUDUL", "TITLE"],
        "Abstrak": ["Abstrak", "Abstract", "ABSTRAK", "ABSTRACT"],
        "Pendahuluan": ["INTRODUCTION", "PENDAHULUAN"],
        "Metodologi": ["RESEARCH METHODOLOGY", "METODOLOGI PENELITIAN", "RESEARCH METHOD", "METODE PENELITIAN"],
        "Hasil": ["RESULT AND DISCUSSION", "HASIL DAN PEMBAHASAN"],
        "Kesimpulan": ["CONCLUSION", "KESIMPULAN"],
        "Referensi": ["REFERENCES", "REFERENSI"]
    }

    for section_name, feedback in review_data.items():
        keywords = sections.get(section_name, [])
        para_idx = find_section_paragraph(doc, keywords)
        
        if para_idx is not None:
            add_comment_to_paragraph(doc.paragraphs[para_idx], feedback)
        else:
            # Jika tidak ketemu, coba tambah ke paragraph pertama
            if doc.paragraphs:
                add_comment_to_paragraph(doc.paragraphs[0], f"[{section_name}]: {feedback}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_text_from_docx(file):
    doc = Document(BytesIO(file.getvalue()))
    return "\n".join([para.text for para in doc.paragraphs])

def get_ai_review_structured(text, api_key, role_choice, language_choice, level):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-3-flash-preview') 
    
    prompt = f"""
    Bertindaklah sebagai {role_choice}. Bahasa: {language_choice}. Level Kritik: {level}.
    Berikan review manuskrip dalam format JSON murni. 
    PENTING: Gunakan tanda kutip tunggal (') di dalam teks review, BUKAN kutip ganda (").
    Format JSON:
    {{
      "Judul": "...", "Abstrak": "...", "Pendahuluan": "...", 
      "Metodologi": "...", "Hasil": "...", "Kesimpulan": "...", "Referensi": "..."
    }}
    Manuskrip: {text}
    """
    
    response = model.generate_content(prompt)
    clean_json = response.text.replace('```json', '').replace('```', '').strip()
    
    # Logic slicing JSON untuk keamanan parsing
    start_idx = clean_json.find('{')
    end_idx = clean_json.rfind('}')
    return json.loads(clean_json[start_idx:end_idx + 1])

# --- 3. KONFIGURASI STREAMLIT ---

load_dotenv()
saved_api_key = os.getenv("GEMINI_API_KEY")

st.set_page_config(page_title="AI Reviewer 2026", page_icon="📝", layout="wide")
st.title("📝 AI Manuscript Reviewer")

with st.sidebar:
    st.header("⚙️ Pengaturan")
    role = st.selectbox("Peran Reviewer:", ["Reviewer 1 (Substansi)", "Reviewer 2 (Teknis)"])
    output_lang = st.selectbox("Bahasa Hasil:", ["Bahasa Indonesia", "English"])
    critique_level = st.select_slider("Level Kritik:", options=["Cukup Kritis", "Kritis", "Sangat Kritis"], value="Kritis")
    api_key = saved_api_key if saved_api_key else st.text_input("Gemini API Key", type="password")

uploaded_file = st.file_uploader("Pilih file Word (.docx)", type=["docx"])

if uploaded_file:
    if st.button("🚀 Mulai Analisis AI"):
        if not api_key:
            st.error("Masukkan API Key!")
        else:
            with st.spinner("Menghubungi AI..."):
                try:
                    text_content = extract_text_from_docx(uploaded_file)
                    review_dict = get_ai_review_structured(text_content, api_key, role, output_lang, critique_level)
                    st.session_state['review_dict'] = review_dict
                    st.success("Analisis selesai!")
                except Exception as e:
                    st.error(f"Gagal: {e}")

if 'review_dict' in st.session_state:
    st.divider()
    st.subheader("📋 Pilih Komentar")
    current_selection = {}
    cols = st.columns(2)
    items = list(st.session_state['review_dict'].items())
    
    for i, (section, feedback) in enumerate(items):
        with cols[i % 2]:
            if st.checkbox(f"**{section}**", key=f"chk_{section}"):
                st.info(feedback)
                current_selection[section] = feedback
    
    if current_selection:
        if st.button("💾 Generate Word"):
            with st.spinner("Memproses file..."):
                final_docx = process_word_with_comments(uploaded_file, current_selection)
                st.download_button("📥 Unduh Hasil", data=final_docx, file_name=f"REVIEW_{uploaded_file.name}")
